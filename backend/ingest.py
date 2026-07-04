# -*- coding: utf-8 -*-
"""Ингест документов: PDF/DOCX -> нормализованный текст + метаданные.

Выход: data/cache/docs.jsonl (одна строка = один документ)
  {id, filename, ext, title, year, lang, geo_hint, n_chars, text_path}
и data/cache/txt/<id>.txt с полным текстом.
"""
import json
import os
import re
import sys
import hashlib
import traceback

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langdetect import detect, DetectorFactory

DetectorFactory.seed = 0

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RAW = os.path.join(ROOT, "data", "raw", "obzory")
CACHE = os.path.join(ROOT, "data", "cache")
TXT = os.path.join(CACHE, "txt")
os.makedirs(TXT, exist_ok=True)

YEAR_RE = re.compile(r"(19[89]\d|20[0-2]\d)")

# Простая эвристика географии по упоминаниям (уточняется LLM-слоем позже)
GEO_FOREIGN = re.compile(
    r"(зарубежн|мировой практик|за рубежом|Австрали|Канад|Чили|Финлянд|Китa|Китай|США|"
    r"Outotec|Glencore|Vale|BHP|Boliden|Sherritt|Jinchuan|Sumitomo)", re.I)
GEO_RU = re.compile(
    r"(Норильск|Кольск|Мончегорск|Надеждинск|Талнах|ГМК|отечествен|Россия|российск|"
    r"Урал|Гипроникель)", re.I)


def extract_pdf(path: str) -> str:
    doc = fitz.open(path)
    parts = []
    for page in doc:
        parts.append(page.get_text("text"))
    doc.close()
    return "\n".join(parts)


def extract_docx(path: str) -> str:
    d = DocxDocument(path)
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def guess_year(filename: str, text: str) -> int | None:
    # приоритет: имя файла, затем первые 3000 символов текста
    m = YEAR_RE.findall(filename)
    if m:
        return int(m[-1])
    m = YEAR_RE.findall(text[:3000])
    if m:
        # берём самый частый год в шапке документа
        from collections import Counter
        return int(Counter(m).most_common(1)[0][0])
    return None


def guess_lang(text: str) -> str:
    sample = text[:2000].strip()
    if not sample:
        return "unknown"
    try:
        code = detect(sample)
        return code if code in ("ru", "en") else code
    except Exception:
        return "unknown"


def geo_hint(text: str) -> str:
    f = len(GEO_FOREIGN.findall(text))
    r = len(GEO_RU.findall(text))
    if f and r:
        return "mixed"
    if f:
        return "foreign"
    if r:
        return "ru"
    return "unknown"


def main() -> None:
    out_path = os.path.join(CACHE, "docs.jsonl")
    done_ids = set()
    if os.path.exists(out_path):
        with open(out_path, encoding="utf-8") as f:
            for line in f:
                try:
                    done_ids.add(json.loads(line)["id"])
                except Exception:
                    pass

    files = sorted(os.listdir(RAW))
    ok, skip, fail = 0, 0, 0
    with open(out_path, "a", encoding="utf-8") as out:
        for fn in files:
            path = os.path.join(RAW, fn)
            ext = fn.rsplit(".", 1)[-1].lower()
            doc_id = hashlib.md5(fn.encode("utf-8")).hexdigest()[:12]
            if doc_id in done_ids:
                skip += 1
                continue
            if ext not in ("pdf", "docx"):
                skip += 1
                continue
            try:
                text = extract_pdf(path) if ext == "pdf" else extract_docx(path)
                text = re.sub(r"[ \t]+", " ", text)
                text = re.sub(r"\n{3,}", "\n\n", text).strip()
                if len(text) < 200:
                    print(f"SKIP(short) {fn}", flush=True)
                    skip += 1
                    continue
                txt_path = os.path.join(TXT, doc_id + ".txt")
                with open(txt_path, "w", encoding="utf-8") as tf:
                    tf.write(text)
                rec = {
                    "id": doc_id,
                    "filename": fn,
                    "ext": ext,
                    "title": os.path.splitext(fn)[0],
                    "year": guess_year(fn, text),
                    "lang": guess_lang(text),
                    "geo_hint": geo_hint(text),
                    "n_chars": len(text),
                    "text_path": os.path.relpath(txt_path, ROOT),
                }
                out.write(json.dumps(rec, ensure_ascii=False) + "\n")
                out.flush()
                ok += 1
                print(f"OK {fn} ({len(text)} chars, {rec['lang']}, {rec['year']})", flush=True)
            except Exception:
                fail += 1
                print(f"FAIL {fn}\n{traceback.format_exc()}", flush=True)
    print(f"\nDONE ok={ok} skip={skip} fail={fail}", flush=True)


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    main()
